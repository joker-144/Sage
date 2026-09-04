"""
PaperExport — 论文成稿导出（LaTeX / Word）

提供 markdown → LaTeX / Word 转换：
  - to_latex: 纯文本转换，产出可编译的 article + ctex 文档
  - to_docx: 需要 python-docx（可选依赖 [paper]）

支持的块级元素：标题、段落、无序列表、**管道表格**（学术论文三线表）、
**图片**（`![alt](path)`）。表格转换规则：
  - LaTeX  → booktabs 三线表（\\toprule/\\midrule/\\bottomrule）+ table 浮动环境
  - Word   → python-docx 原生表格（表头加粗）
  - 表格前一行若为 "表N：标题" / "Table N: xxx" 形式，自动作为表题（caption）

导出的 markdown 通常来自 PaperProject.read_draft()（最终 paper.md）。
"""
from __future__ import annotations

import re
from pathlib import Path

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_CITE_RE = re.compile(r"\[CITE:\s*([^\]]*)\]")
_IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
# 表题：表格紧邻的上一行（表3：xxx / Table 3: xxx / 表 3 xxx）
_TABLE_CAPTION_RE = re.compile(r"^(表\s*\d*|Table\s*\d*)\s*[:：]?\s*(.*)$", re.IGNORECASE)


def _inline_docx(md: str) -> str:
    """行内元素转换（Word）：去掉 markdown 标记，保留纯文本与引用编号。"""
    md = re.sub(r"\*\*(.+?)\*\*", r"\1", md)
    md = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", md)
    return md


def _escape_latex(text: str) -> str:
    """LaTeX 转义 + 行内元素转换（用于标题/段落/表格单元格）。

    顺序关键：先转义特殊字符（& % # _ 等），再做 **加粗** / *斜体* /
    [CITE:] → \\cite 转换，保证生成的命令本身不被二次转义。
    """
    text = text or ""
    for ch in ("\\", "{", "}", "&", "%", "$", "#", "_", "~", "^"):
        text = text.replace(ch, "\\" + ch)
    text = re.sub(r"\*\*(.+?)\*\*", r"\\textbf{\1}", text)
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\\textit{\1}", text)
    text = _CITE_RE.sub(r"\\cite{\1}", text)
    return text


def _split_table_row(line: str) -> list[str]:
    """拆分管道表格行为单元格（去掉首尾管道，按 | 分列，容忍转义 \\|）"""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    cells: list[str] = []
    buf = ""
    escaped = False
    for ch in line:
        if escaped:
            buf += ch
            escaped = False
        elif ch == "\\":
            buf += ch
            escaped = True
        elif ch == "|":
            cells.append(buf.strip())
            buf = ""
        else:
            buf += ch
    cells.append(buf.strip())
    return cells


def _is_table_separator(line: str) -> bool:
    """判断是否为表格分隔行 |---|---|"""
    stripped = line.strip()
    if "|" not in stripped:
        return False
    cells = _split_table_row(stripped)
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", c) for c in cells)


def _iter_blocks(draft: str) -> list[tuple[str, object]]:
    """把 markdown 拆为 (kind, payload) 块。

    kind 与 payload:
      h1/h2/h3      → str（标题文本）
      para          → str（段落文本）
      bullet        → str（列表项文本）
      table         → dict(caption, headers, rows)
      image         → dict(path, alt)
    """
    blocks: list[tuple[str, object]] = []
    lines = (draft or "").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        m = _HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            kind = "h1" if level == 1 else ("h2" if level == 2 else "h3")
            blocks.append((kind, m.group(2).strip()))
            i += 1
            continue
        if line.strip().startswith("- "):
            blocks.append(("bullet", line.strip()[2:]))
            i += 1
            continue
        img = _IMAGE_RE.match(line.strip())
        if img:
            blocks.append(("image", {"path": img.group(2).strip(), "alt": img.group(1).strip()}))
            i += 1
            continue
        # 管道表格：当前行为表头（含 |），下一行为分隔行
        if "|" in line and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            headers = _split_table_row(line)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append(_split_table_row(lines[i]))
                i += 1
            # 表题：取表格上方最近的一个非空块（若是段落且形如 表N：xxx）
            caption = ""
            for j in range(len(blocks) - 1, -1, -1):
                kind, payload = blocks[j]
                if kind == "para":
                    cm = _TABLE_CAPTION_RE.match(str(payload).strip())
                    if cm:
                        caption = cm.group(2).strip() or str(payload).strip()
                        blocks.pop(j)
                break
            blocks.append(("table", {"caption": caption, "headers": headers, "rows": rows}))
            continue
        if line.strip():
            blocks.append(("para", line.strip()))
        i += 1
    return blocks


def to_latex(draft: str, title: str = "") -> str:
    """把 markdown 草稿转换为可编译的 LaTeX 文档（article + ctex 中文支持）。

    表格转 booktabs 三线表（浮动 table 环境），图片转 figure 环境（graphicx）。
    """
    body: list[str] = []
    in_list = False

    def _close_list():
        nonlocal in_list
        if in_list:
            body.append("\\end{itemize}")
            in_list = False

    for kind, payload in _iter_blocks(draft):
        if kind in ("h1", "h2", "h3"):
            _close_list()
            heading = _escape_latex(str(payload))
            cmd = {"h1": "\\section", "h2": "\\subsection"}.get(kind, "\\subsubsection")
            body.append(f"{cmd}{{{heading}}}")
        elif kind == "bullet":
            if not in_list:
                body.append("\\begin{itemize}")
                in_list = True
            body.append("  \\item " + _escape_latex(str(payload)))
        elif kind == "table":
            _close_list()
            tbl = payload
            body.append("\\begin{table}[htbp]")
            body.append("\\centering")
            if tbl["caption"]:
                body.append(f"\\caption{{{_escape_latex(tbl['caption'])}}}")
            col_spec = "c" * max(len(tbl["headers"]), 1)
            body.append(f"\\begin{{tabular}}{{{col_spec}}}")
            body.append("\\toprule")
            body.append(" & ".join(_escape_latex(h) for h in tbl["headers"]) + " \\\\")
            body.append("\\midrule")
            for row in tbl["rows"]:
                # 列数不齐时以表头列数为准补空
                cells = list(row) + [""] * (len(tbl["headers"]) - len(row))
                body.append(" & ".join(_escape_latex(c) for c in cells[: len(tbl["headers"])]) + " \\\\")
            body.append("\\bottomrule")
            body.append("\\end{tabular}")
            body.append("\\end{table}")
        elif kind == "image":
            _close_list()
            img = payload
            body.append("\\begin{figure}[htbp]")
            body.append("\\centering")
            body.append(f"\\includegraphics[width=0.8\\textwidth]{{{img['path']}}}")
            if img["alt"]:
                body.append(f"\\caption{{{_escape_latex(img['alt'])}}}")
            body.append("\\end{figure}")
        else:
            _close_list()
            body.append(_escape_latex(str(payload)))
    _close_list()

    doc_title = _escape_latex(title or "论文")
    return (
        "\\documentclass{article}\n"
        "\\usepackage[UTF8]{ctex}\n"
        "\\usepackage{booktabs}\n"
        "\\usepackage{graphicx}\n"
        f"\\title{{{doc_title}}}\n"
        "\\author{}\n"
        "\\begin{document}\n"
        "\\maketitle\n\n"
        + "\n".join(body)
        + "\n\n\\end{document}\n"
    )


def to_docx(draft: str, out_path) -> Path:
    """把 markdown 草稿导出为 Word 文档（需要 python-docx 可选依赖）。

    表格生成原生 Word 表格（表头加粗），图片在文件存在时内嵌、
    不存在时插入占位段落。
    """
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "Word 导出需要 python-docx，请安装: pip install 'sage-paper[paper]'"
        ) from e
    doc = Document()
    for kind, payload in _iter_blocks(draft):
        if kind in ("h1", "h2", "h3"):
            doc.add_heading(str(payload), level=int(kind[1]))
        elif kind == "bullet":
            doc.add_paragraph(_inline_docx(str(payload)), style="List Bullet")
        elif kind == "table":
            tbl = payload
            headers, rows = tbl["headers"], tbl["rows"]
            table = doc.add_table(rows=1 + len(rows), cols=max(len(headers), 1))
            table.style = "Table Grid"
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = _inline_docx(h)
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
            for r, row in enumerate(rows, start=1):
                cells = list(row) + [""] * (len(headers) - len(row))
                for j in range(len(headers)):
                    table.rows[r].cells[j].text = _inline_docx(cells[j])
            if tbl["caption"]:
                doc.add_paragraph(_inline_docx(tbl["caption"]))
        elif kind == "image":
            img = payload
            img_path = Path(img["path"])
            embedded = False
            if img_path.is_file():
                try:
                    doc.add_picture(str(img_path))
                    embedded = True
                except Exception:
                    embedded = False
            if embedded and img["alt"]:
                doc.add_paragraph(_inline_docx(img["alt"]))
            elif not embedded:
                doc.add_paragraph(f"[图: {img['alt'] or '未命名'} ({img['path']})]")
        else:
            doc.add_paragraph(_inline_docx(str(payload)))
    out = Path(out_path)
    doc.save(str(out))
    return out
